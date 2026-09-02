from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


TITLE = (
    "FixProof: A Validation-Oriented Agentic Security System for "
    "AI-Generated Vulnerability Remediation"
)


COMPLETED_TASKS = (
    "Approximately 30 hours total (student must confirm against actual time): "
    "(1) about 8 hours refining the research question, reviewing background "
    "literature and course expectations, narrowing the scope to a controlled "
    "JavaScript/Express prototype, and documenting trust boundaries and the "
    "artifact architecture; (2) about 14 hours implementing and exercising the "
    "scanner/provenance pipeline, focused prompt construction, OpenAI remediation "
    "agent, isolated patch workspaces, syntax/SAST comparison, CWE-specific "
    "security tests, functional regression tests, retry evidence, and the "
    "deterministic decision policy; and (3) about 8 hours implementing evidence-"
    "bound human adjudication, manifest-selected evaluation metrics, the read-only "
    "dashboard, reproducibility checks, a multi-scenario demo runner, API-key "
    "handling, documentation, and regression testing. The current reproducibility "
    "gate passes 37 automated tests."
)


TIMELINE_ROWS = [
    ("W1 (Aug 25-Aug 31)", "Refine the problem statement, research questions, novelty claim, and bounded CS6727 scope using proposal feedback and primary sources.", "Completed"),
    ("W1", "Establish the FixProof architecture, three controlled benchmark cases, evidence schema, initial validation results, and reproducibility baseline.", "Completed"),
    ("W2 (Sep 1-Sep 7)", "Build a literature-evidence matrix covering AI code security, secure development validation, and relevant CWE definitions.", "Planned"),
    ("W2", "Freeze the threat model, independent ground-truth procedure, benchmark inclusion criteria, and artifact schema for the next experiment revision.", "Planned"),
    ("W3 (Sep 8-Sep 14)", "Automate a disposable end-to-end live experiment runner with explicit attempt metadata and no authoritative-artifact overwrite.", "Planned"),
    ("W3", "Expand negative tests for malformed artifacts, stale hashes, identity mismatches, occupied ports, and inconclusive validators.", "Planned"),
    ("W4 (Sep 15-Sep 21)", "Run pilot repeated remediation trials for CWE-79, CWE-89, and CWE-22 using fixed prompts and a recorded model configuration.", "Planned"),
    ("W4", "Audit pilot outputs, revise only documented protocol defects, and lock the main evaluation procedure before collecting final observations.", "Planned"),
    ("W5 (Sep 22-Sep 28)", "Implement a SAST-only comparison/ablation so the evaluation can measure what independent runtime and functional evidence adds.", "Planned"),
    ("W5", "Define repetition counts, model/configuration recording rules, exclusion rules, and descriptive-analysis limits for the controlled sample.", "Planned"),
    ("W6 (Sep 29-Oct 5)", "Execute the frozen main remediation runs across the three benchmark vulnerability classes and retain unchanged model responses.", "Planned"),
    ("W6", "Record execution metadata needed for auditability, including response IDs, model configuration, attempt lineage, duration, and API usage when available.", "Planned"),
    ("W7 (Oct 6-Oct 12)", "Run syntax, SAST, targeted security, and functional validation for every selected attempt and generate deterministic decisions.", "Planned"),
    ("W7", "Review evidence-bound human packets for disagreement cases and record reviewer conclusions separately from automated decisions.", "Planned"),
    ("W8 (Oct 13-Oct 19)", "Exercise failure branches through controlled malformed, incomplete, false-success, and validator-inconclusive test inputs.", "Planned"),
    ("W8", "Complete artifact provenance, SHA-256 bindings, metric-scope separation, and clean-environment reproduction instructions.", "Planned"),
    ("W9 (Oct 20-Oct 26)", "Analyze remediation outcomes, SAST/runtime disagreements, functional regressions, and retry improvement without generalizing beyond the sample.", "Planned"),
    ("W9", "Draft the architecture, threat model, methodology, and implementation sections using evidence from the frozen repository.", "Planned"),
    ("W10 (Oct 27-Nov 2)", "Refine the dashboard and controlled CLI demonstration to explain evidence channels, decisions, and human boundaries clearly.", "Planned"),
    ("W10", "Perform a clean-environment reproducibility rehearsal and document every setup failure or platform-specific limitation.", "Planned"),
    ("W11 (Nov 3-Nov 9)", "Freeze the authoritative experiment manifest and regenerate the machine-readable and human-readable evaluation reports.", "Planned"),
    ("W11", "Draft the results section with attempt-level tables, aggregate descriptive metrics, retry analysis, and control separation.", "Planned"),
    ("W12 (Nov 10-Nov 16)", "Draft the discussion, limitations, threats-to-validity, deployment-feasibility, and ethical/privacy sections.", "Planned"),
    ("W12", "Audit citations, sanitized prompt documentation, AI-assistance disclosure, and removal of credentials or confidential information.", "Planned"),
    ("W13 (Nov 17-Nov 23)", "Create the final presentation narrative, architecture visual, evaluation figures, and bounded live demonstration sequence.", "Planned"),
    ("W13", "Rehearse the demo, measure timing, and prepare a recorded-artifact fallback that preserves the same evidence and decisions.", "Planned"),
    ("W14 (Nov 24-Nov 30)", "Incorporate instructional and peer feedback into the implementation explanation, evaluation interpretation, and limitations.", "Planned"),
    ("W14", "Complete a code, test, documentation, and artifact-integrity review; close only issues that affect the stated research claims.", "Planned"),
    ("W15 (Dec 1-Dec 7)", "Finalize the report, appendices, figures, references, reproducibility instructions, and generative-AI disclosure.", "Planned"),
    ("W15", "Run the final verification gate, check hashes and dashboard data, review similarity results, and proofread the submission package.", "Planned"),
    ("W16 (Dec 8-Dec 14)", "Package the sanitized repository snapshot, experiment evidence, demo instructions, and final presentation materials.", "Planned"),
    ("W16", "Verify Canvas requirements and due dates, remove secrets and temporary data, and submit the final deliverables.", "Planned"),
]


REFERENCES = """[1] H. Pearce, B. Ahmad, B. Tan, B. Dolan-Gavitt, and R. Karri, “Asleep at the Keyboard? Assessing the Security of GitHub Copilot’s Code Contributions,” 2022 IEEE Symposium on Security and Privacy, pp. 754-768, 2022, doi: 10.1109/SP46214.2022.9833571.
[2] N. Perry, M. Srivastava, D. Kumar, and D. Boneh, “Do Users Write More Insecure Code with AI Assistants?” Proceedings of the 2023 ACM SIGSAC Conference on Computer and Communications Security, pp. 2785-2799, 2023, doi: 10.1145/3576915.3623157.
[3] G. Sandoval, H. Pearce, T. Nys, R. Karri, S. Garg, and B. Dolan-Gavitt, “Lost at C: A User Study on the Security Implications of Large Language Model Code Assistants,” 32nd USENIX Security Symposium, pp. 2205-2222, 2023.
[4] M. Souppaya, K. Scarfone, and D. Dodson, Secure Software Development Framework (SSDF) Version 1.1, NIST SP 800-218, Feb. 2022, doi: 10.6028/NIST.SP.800-218.
[5] MITRE, “CWE-79: Improper Neutralization of Input During Web Page Generation (‘Cross-site Scripting’),” Common Weakness Enumeration, accessed Aug. 31, 2026, https://cwe.mitre.org/data/definitions/79.html.
[6] MITRE, “CWE-89: Improper Neutralization of Special Elements used in an SQL Command (‘SQL Injection’),” Common Weakness Enumeration, accessed Aug. 31, 2026, https://cwe.mitre.org/data/definitions/89.html.
[7] MITRE, “CWE-22: Improper Limitation of a Pathname to a Restricted Directory (‘Path Traversal’),” Common Weakness Enumeration, accessed Aug. 31, 2026, https://cwe.mitre.org/data/definitions/22.html.
[8] OpenAI, “GPT-5.2 Model,” OpenAI API Documentation, accessed Aug. 31, 2026, https://developers.openai.com/api/docs/models/gpt-5.2."""


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def set_title(paragraph, text: str, size: int) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_document(template: Path, output: Path) -> None:
    document = Document(template)
    paragraphs = document.paragraphs

    set_text(paragraphs[0], "Section: CS")
    set_title(paragraphs[1], TITLE, 16)
    set_title(paragraphs[2], "Tony Tran", 12)
    period = paragraphs[3].insert_paragraph_before(
        "Progress Report I — Draft prepared August 31, 2026"
    )
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        paragraphs[4],
        "AI coding assistants can accelerate software development, but generated "
        "code and generated vulnerability fixes may be insecure, functionally "
        "incorrect, or over-trusted. Pearce et al. found that approximately 40% "
        "of 1,689 Copilot-generated programs in their security scenarios were "
        "vulnerable, while Perry et al. found that AI-assisted participants wrote "
        "less secure code and were more likely to believe their code was secure. "
        "Other studies report more context-dependent effects, which reinforces the "
        "need to evaluate concrete outputs rather than assume that AI assistance is "
        "always harmful or always safe. In vulnerability remediation, a static "
        "analysis finding disappearing is not sufficient proof: a candidate can "
        "retain exploitable behavior, create a new weakness, or break intended "
        "functionality. The practical research problem is the lack of a small, "
        "auditable workflow that independently validates AI-generated security "
        "patches before a human is asked to accept them."
    )
    set_text(
        paragraphs[5],
        "FixProof studies this problem through three controlled JavaScript/Express "
        "benchmarks—reflected XSS (CWE-79), SQL injection (CWE-89), and path "
        "traversal (CWE-22). The novelty is the evidence-oriented orchestration: "
        "scanner provenance, focused remediation context, independent runtime and "
        "functional evidence, deterministic policy outcomes, immutable artifact "
        "bindings, and a separate human-review record are treated as one research "
        "system rather than allowing the remediation model or one scanner result "
        "to determine success."
    )

    set_text(
        paragraphs[7],
        "This research will design, implement, and evaluate FixProof, a validation-"
        "oriented agentic security prototype, in order to determine whether combined "
        "SAST, targeted runtime security, and functional evidence can prevent false "
        "confidence in AI-generated vulnerability remediation. FixProof scans a "
        "controlled benchmark, normalizes and correlates findings while preserving "
        "rule origin, constructs a focused structured prompt, records one model "
        "candidate, applies it only in an isolated workspace, and independently "
        "runs syntax, candidate-SAST, exploit-oriented, and regression checks. A "
        "deterministic policy then produces REJECT, READY_FOR_HUMAN_REVIEW, or "
        "NEEDS_HUMAN_ADJUDICATION; any human conclusion is stored separately."
    )
    set_text(
        paragraphs[8],
        "Deliverables are a documented Python/JavaScript research prototype, three "
        "controlled vulnerable applications and ground-truth tests, structured "
        "experiment artifacts, a reproducibility/test gate, evaluation metrics and "
        "reports, an evidence dashboard, a bounded CLI demonstration, and a final "
        "paper discussing efficacy and limitations. The semester scope excludes "
        "general-purpose SAST/DAST replacement, autonomous deployment, production "
        "approval, arbitrary hostile-repository execution, and broad claims about "
        "all models or programming languages."
    )

    set_text(
        paragraphs[10],
        "Research, scope, and architecture (approximately 8 hours): refined the "
        "primary question and secondary retry question; reviewed security literature "
        "and the CS6727 requirements; bounded the study to three Express benchmarks; "
        "and documented the two-plane remediation/validation design, trust boundaries, "
        "artifact flow, decision policy, and limitations."
    )
    set_text(
        paragraphs[11],
        "Implementation and validation (approximately 14 hours): implemented or "
        "integrated Semgrep scanning, controlled-rule labels and provenance, finding "
        "normalization/correlation, focused contexts and prompts, structured OpenAI "
        "remediation, isolated patch workspaces, syntax and semantic SAST comparison, "
        "CWE-79/CWE-89/CWE-22 runtime security tests, functional regression tests, "
        "retry feedback, and deterministic evidence-aware dispositions."
    )
    set_text(
        paragraphs[12],
        "Evaluation, human boundary, and reproducibility (approximately 8 hours): "
        "created immutable adjudication packets and reviewer results, an authoritative "
        "manifest, SHA-256 artifact bindings, primary/control metric separation, a "
        "false-success control, generated reports, a read-only dashboard, local API-"
        "key handling, a multi-case demonstration runner, and regression tests. The "
        "current gate passes 37 tests and reports three cases, four AI attempts, one "
        "separated non-AI control, and two completed required adjudications. The three "
        "completed-task categories total approximately 30 hours; these allocations "
        "must be reconciled with my actual time log before submission."
    )

    set_text(
        paragraphs[15],
        "Over the next two weeks I will (1) create a structured related-work matrix "
        "and tighten the novelty claim; (2) freeze the threat model, benchmark "
        "ground truth, artifact schema, and evaluation protocol; (3) automate a "
        "disposable end-to-end live experiment path so fresh model outputs can be "
        "validated without entering the authoritative dataset; and (4) expand "
        "negative tests for malformed/stale artifacts, evidence conflicts, occupied "
        "ports, and inconclusive validators."
    )
    set_text(
        paragraphs[16],
        "I will then run pilot repeated trials across the three CWEs using a fixed "
        "prompt/model configuration, examine whether the procedure needs correction, "
        "and lock the main evaluation before collecting final observations. I will "
        "also design a SAST-only comparison/ablation so the final evaluation can "
        "measure the value added by runtime, functional, policy, and human-review "
        "layers."
    )

    set_text(paragraphs[17], "Questions I have or issues I am running into:")
    set_text(
        paragraphs[18],
        "The current system is technically complete enough for a pilot, but the "
        "sample is intentionally small (three CWEs, JavaScript/Express, one primary "
        "model configuration). This limits statistical and production generalization. "
        "Semgrep CE missed the SQLi benchmark until a clearly labeled, benchmark-"
        "focused controlled rule was added; XSS and path-traversal candidates also "
        "produce SAST/runtime disagreements. The XSS validator inspects reflected "
        "output but does not execute JavaScript in a browser, and filesystem copies "
        "are experiment isolation rather than hardened containers. These limitations "
        "will remain explicit rather than being hidden through additional features."
    )
    set_text(
        paragraphs[19],
        "Guidance requested from the instructional team: Is a controlled case-study "
        "evaluation with repeated attempts and a SAST-only ablation appropriate for "
        "the expected practicum rigor, or is a larger benchmark sample expected? "
        "Please also confirm whether a sanitized appendix summarizing actual AI "
        "prompts/tools is sufficient for the course disclosure requirement or whether "
        "full prompt transcripts should be submitted separately."
    )

    set_text(
        paragraphs[21],
        "FixProof uses a design-science and controlled experimental methodology. For "
        "each benchmark, I first establish vulnerability ground truth independently "
        "of the scanner; run standard and explicitly labeled controlled rules; "
        "normalize/correlate evidence with provenance; and provide the remediation "
        "model only the canonical finding and focused source context, withholding "
        "ground-truth test answers. The unchanged model response is applied to an "
        "isolated copy. Independent stages then check syntax, rescan and semantically "
        "compare SAST findings, execute vulnerability-specific payloads, and verify "
        "intended behavior. A deterministic policy converts those channels into a "
        "disposition, and disagreement cases receive a separate evidence-bound human "
        "record."
    )
    set_text(
        paragraphs[22],
        "The experimental unit is one remediation attempt. Primary metrics include "
        "target-SAST resolution, targeted security pass, functional preservation, new "
        "finding rate, SAST false success, SAST/runtime disagreement, retry improvement, "
        "and adjudication completion. Non-AI policy-coverage controls are labeled and "
        "excluded from AI-attempt denominators. A manifest selects authoritative "
        "artifacts and the report builder verifies identities, attempt lineage, hashes, "
        "evidence consistency, and metric scope before aggregation."
    )
    set_text(
        paragraphs[23],
        "The final analysis will compare the full FixProof policy with a SAST-only "
        "interpretation, use repeated fixed-protocol trials where feasible, report "
        "attempt-level results and descriptive rates, and avoid population-level "
        "claims. Reproduction uses a pinned project environment, generated JSON/"
        "Markdown reports, 37 current automated tests, and a read-only dashboard."
    )
    set_text(paragraphs[24], "")
    for index in range(25, 30):
        set_text(paragraphs[index], "")

    set_text(
        paragraphs[31],
        "The timeline below contains only actionable FixProof research, implementation, "
        "evaluation, writing, and delivery tasks. Exact milestone dates must be checked "
        "against Canvas before submission."
    )

    table = document.tables[0]
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for week, task, status in TIMELINE_ROWS:
        cells = table.add_row().cells
        cells[0].text = week
        cells[1].text = task
        cells[2].text = status

    set_text(
        paragraphs[34],
        "Current pilot/engineering results: the authoritative matrix contains four "
        "AI attempts across three cases and one separated deterministic non-AI "
        "false-success control. The current results are target-SAST resolution 1/4 "
        "(25%), targeted security pass 4/4 (100%), functional preservation 3/4 "
        "(75%), new-finding rate 0/4, retry improvement 1/1, two SAST/runtime "
        "disagreements, and adjudication completion 2/2. These values describe only "
        "the controlled pilot and are not production performance estimates. The final "
        "evaluation will repeat a frozen protocol, compare full FixProof decisions "
        "against a SAST-only interpretation, exercise failure branches, verify "
        "artifact integrity and reproducibility, and analyze security/functionality "
        "tradeoffs and threats to validity."
    )
    set_text(
        paragraphs[36],
        "Provisional final-report outline: 1. Abstract; 2. Introduction and problem "
        "statement; 3. Related work; 4. Research questions, requirements, threat model, "
        "and scope; 5. FixProof architecture and trust boundaries; 6. Experimental "
        "methodology and benchmark ground truth; 7. Implementation; 8. Evaluation "
        "results and SAST-only comparison; 9. Discussion, human adjudication, and "
        "deployment feasibility; 10. Limitations and threats to validity; 11. Ethics, "
        "privacy, and generative-AI disclosure; 12. Conclusion; Appendices: artifact "
        "schemas, prompt documentation, reproducibility commands, and detailed tables."
    )
    set_text(paragraphs[38], REFERENCES)
    set_text(
        paragraphs[39],
        "Repository evidence reviewed for this report: README.md, "
        "IMPLEMENTATION_GUIDE.md, docs/architecture.md, docs/methodology.md, "
        "docs/evaluation-report.md, the authoritative experiment manifest, test "
        "suite, benchmark applications, and bound validation/adjudication artifacts."
    )

    set_text(
        paragraphs[41],
        "Generative-AI assistance disclosure: OpenAI ChatGPT/Codex was used as a "
        "supplement for brainstorming, code scaffolding, debugging/test orchestration, "
        "documentation editing, and preparation of this progress-report draft. The "
        "OpenAI API model recorded as gpt-5.2 generated the candidate remediation "
        "artifacts evaluated by FixProof. Tony Tran remains responsible for the work, "
        "reviewed the outputs, and used independent scanner, runtime, functional, "
        "policy, and human-review evidence rather than treating AI output as ground "
        "truth. Before submission, attach or retain a sanitized log identifying the "
        "actual tools and prompts used; exclude API keys and other credentials."
    )
    set_text(
        paragraphs[42],
        "Privacy and reproducibility note: the experiment uses synthetic local "
        "benchmarks and does not require employer, proprietary, or personal data. API "
        "credentials are loaded from an ignored .env file and are not stored in source "
        "or artifacts. Reproduce the selected report with: $env:PYTHONPATH=\"src\"; "
        ".\\.venv\\Scripts\\python.exe -m fixproof.reproduce --verify. Run the full "
        "controlled demo matrix with: powershell -ExecutionPolicy Bypass -File "
        ".\\demo-test.ps1 -Suite."
    )
    set_text(paragraphs[43], "")

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["List Paragraph"].font.name = "Arial"
    styles["List Paragraph"].font.size = Pt(10)

    document.core_properties.title = TITLE
    document.core_properties.author = "Tony Tran"
    document.core_properties.subject = "CS6727 Progress Report I"
    document.core_properties.comments = (
        "Draft generated from the course template; student review required."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build FixProof Progress Report I.")
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    fill_document(args.template.resolve(), args.output.resolve())
    print(f"Progress report draft: {args.output.resolve()}")


if __name__ == "__main__":
    main()
